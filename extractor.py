import docx2txt
import nltk
import re
from docx import Document
import docx
import PyPDF2
from pdfReader import convert_pdf_to_txt
from pdfFontFinder import font_style
from pdfPageIteratorusingFitz import total_lines_and_char
import locationtagger 

import pandas as pd
import tabula
import os
import spacy
from spacy.matcher import Matcher
from pprint import pprint
import regex as re
from pathlib import Path
import logging.config
from spacy.matcher import PhraseMatcher
from spacy.tokens import Span
from spacy.lang.en import English
from spacy.pipeline import EntityRuler
import re
import nltk
from nltk.corpus import stopwords
import os
import pandas as pd
import spacy
import en_core_web_sm

nltk.download('stopwords')
stop = stopwords.words('english')
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('maxent_ne_chunker')
nltk.download('words')
nltk.download('punkt')
nltk.download('wordnet')
nltk.download('averaged_perceptron_tagger')

nlp = English()
nlp = spacy.load("en_core_web_md")
matcher = Matcher(nlp.vocab)
nlp = en_core_web_sm.load()


def preprocess_data(sen):
    sentences=[]
    sen=[el.strip() for el in sen.split("\n") if len(el) > 0]
    for sentence in sen:
        sentence=nltk.sent_tokenize(sentence)
        for a in sentence:
            sentences.append(a)
    #print(sentences)
    sen=[nltk.word_tokenize(sent) for sent in sentences]
    #print(sentences)
    sen=[nltk.pos_tag(sent) for sent in sen]
    return sen

def cleanText(text):
    # Removing New Lines and New tabs
    text = text.replace('\n', ' ')
    text = text.replace('\t', ' ')
    # Removing Non UTF-8 Characters or Symbols
    text = bytes(text, 'utf-8').decode('utf-8', 'ignore')
    text=re.sub('http\S+\s*', ' ', text)
    text=re.sub('RT|cc', ' ', text)
    text=re.sub('#\S+', '', text)
    text=re.sub('@\S+', '', text)
    text=re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', text)
    text=re.sub('\s+', ' ', text)
    text=re.sub(r'[^\x00-\x7f]', r' ', text)
    return text

def spacyProcessText(text):
    text_cleaned = cleanText(text)
    doc = nlp(text_cleaned)
    # tokens = [(token.text, token.pos_) for token in doc]
    # sentences = [sent for sent in doc.sents]
    # emails = [token.text for token in doc if token.like_email]
    # urls = [token.text for token in doc if token.like_url]
    # ents = [(e.text, e.label_) for e in doc.ents]
    # for ent in doc.ents:
    #     if ent.label_ in ['GPE', 'LOC']:
    #         print(ent.text, ent.start_char, ent.end_char, ent.label_)
    # print(ents)
    # print(emails)
    # print(urls)
    # print(tokens)
    # pprint(sentences)
    return doc

def extract_full_name(nlp_doc):
     pattern = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
     matcher.add("FULL_NAME",  [pattern], on_match=None)
     matches = matcher(nlp_doc)
     for match_id, start, end in matches:
         span = nlp_doc[start:end]
     return span.text

def extract_phone_number(resume_text):
    PHONE_REG = re.compile(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]')
    phone = re.findall(PHONE_REG, resume_text)
    if phone:
        number = ''.join(phone[0])
        if resume_text.find(number) >= 0 and len(number) < 16:
            return number
    return None
    
def get_email_addresses(string):
    r = re.compile(r'[\w\.-]+@[\w\.-]+')
    return r.findall(string)

def extractLocation(spacy_doc):
    for ent in spacy_doc.ents:
        if ent.label_ in ['GPE', 'LOC']:
            return ent.text

#function to extract education
def extract_education(resume_text):
    stop_words = set(nltk.corpus.stopwords.words('english'))
    word_tokens = nltk.tokenize.word_tokenize(resume_text)
    filtered_tokens = [w for w in word_tokens if w not in stop_words]
    filtered_tokens = [w for w in word_tokens if w.isalpha()]
    bigrams_trigrams = list(map(' '.join, nltk.everygrams(filtered_tokens, 2, 3)))
    found_education = set()
    education_db = ["School","school","College","University","college","university","academy","Faculty","faculty","institute","faculdades","Schola","schule","lise","lyceum",                    "lycee","polytechnic","kolej","ünivers","okul","Fculté", "ecole","ingénierie"]
    for token in filtered_tokens:
        if token.lower() in education_db:
            found_education.add(token)
    for ngram in bigrams_trigrams:
        if ngram.lower() in education_db:
            found_education.add(ngram)
    return found_education
    # check for one-grams
    for token in tokens:
        if token.lower() in education:
            educationset.append(token)
    
    # check for bi-grams and tri-grams
    for token in noun_chunks:
        token = token.text.lower().strip()
        if token in skills:
            educationset.append(token)
    return [i.capitalize() for i in set([i.lower() for i in educationset])]

def cleanup(token, lower = True):
    if lower:
        token = token.lower()
    return token.strip()
#education=extract_education(document)     
# 
def extractName(spacy_doc):
    # First name and Last name are always Proper Nouns
    pattern1 = [{'POS': 'PROPN'}, {'POS': 'PROPN'}]
    pattern2 = [{'POS': 'PROPN'}, {'POS': 'SPACE'}, {'POS': 'PROPN'}]
    matcher.add('NAME', [pattern1, pattern2])
    matches = matcher(spacy_doc)
    # Returning the first match of the above two patterns
    for match_id, start, end in matches:
        # print(start, end)
        span = spacy_doc[start:end]
        name = span.text.title()
        return name       

# TO find which font styles used and which font sizes used in a document
def find_font(path):
    ext=path.rsplit('.', 1)[1].lower()
    #print(ext)
    if(ext=="docx"):
        document = Document(path)
        fname=[]
        size=[]
        for para in document.paragraphs:
            for run in para.runs:
                #print(run.font.name)
                #print(run.text)
                if(not(run.font.name==None)):
                    if(not(run.font.name in fname)):
                        fname.append(run.font.name)
                if(not(run.font.size==None)):
                    f=run.font.size
                    if(not(f.pt in size)):
                        #print(f.pt)
                        size.append(f.pt)
    else:
        # If file is pdf
        fname=font_style(path)
    return fname


#function to extract skills
def extract_skills(resume_text):
    stop_words = set(nltk.corpus.stopwords.words('english'))
    word_tokens = nltk.tokenize.word_tokenize(resume_text)
    filtered_tokens = [w for w in word_tokens if w not in stop_words]
    filtered_tokens = [w for w in word_tokens if w.isalpha()]
    bigrams_trigrams = list(map(' '.join, nltk.everygrams(filtered_tokens, 2, 3)))
    found_skills = set()
    skills_db = ["machine learning",
             "deep learning",
             "nlp",
             "natural language processing",
             "mysql",
             "sql",
             "web",
             "django",
             "computer vision",
              "tensorflow",
             "opencv",
             "mongodb",
             "artificial intelligence",
             "ai",
             "flask",
             "robotics",
             "data structures",
             "python",
             "c++",
             "matlab",
             "css",
             "html",
             "github",
             "php",
             "Django"
             "Angular"
             "Spring Boot",
             "NLP",
             "Deep Learning"]
    for token in filtered_tokens:
        if token.lower() in skills_db:
            found_skills.add(token)
    for ngram in bigrams_trigrams:
        if ngram.lower() in skills_db:
            found_skills.add(ngram)
    return found_skills
    # check for one-grams
    for token in tokens:
        if token.lower() in skills:
            skillset.append(token)
    
    # check for bi-grams and tri-grams
    for token in noun_chunks:
        token = token.text.lower().strip()
        if token in skills:
            skillset.append(token)
    return [i.capitalize() for i in set([i.lower() for i in skillset])]

def cleanup(token, lower = True):
    if lower:
        token = token.lower()
    return token.strip()

# To count no. of tables
def count_tables(path):
    ext=path.rsplit('.', 1)[1].lower()
    #print(ext)
    if(ext=="docx"):
        wordDoc = Document(path)
        count=0
        for table in wordDoc.tables:
            count+=1
            # To extract data from table contents
            #for row in table.rows:
                #for cell in row.cells:
                    #print(cell.text)
        
    else:    # for PDF
        df = tabula.read_pdf(path, pages = 'all', multiple_tables = True)
        #print(df)
        count=0
        for table in df:
            count+=1
    return count
#noOfTables=count_tables(path)
#print(noOfTables)

def export_to_csv(path):
    #path=r"C:/Users/Dagdo/OneDrive/Bureau/Projet Stage/CV.pdf"
    ext=path.rsplit('.', 1)[1].lower()
    if(ext=='docx'):
        document=docx2txt.process(path)
    else:
        document=convert_pdf_to_txt(path)
    #print(document)
    df=pd.DataFrame(columns=["Name","Email","Phone_Number","Location","Textline+Totalchar on each Page","Font_Style","Education","Font_Size","Table_Count","Skills"])
    df.loc[0]=[None,None,None,None,None,None,None,None,None,None]
    a=df.values
    name=extract_full_name(document)
    # for loop
    if(len(name)>0):
        a[0][0]=name[0][0]
    emails=get_email_addresses(document)
    if(len(emails)>0):
        a[0][1]=emails[0]
    numbers=extract_phone_number(document)
    if(len(numbers)>0):
        a[0][2]=numbers[0]
    Location=extractLocation(document)
    if(len(Location)>0):
        a[0][3]=Location[0]
    noOfTables=count_tables(path)
    a[0][7]=noOfTables
    font_name=find_font(path)
    a[0][5]=font_name
    education=extract_education(document)
    a[0][8]=education
    total=textLines_plus_char(path)
    a[0][4]=total
    skills=extract_skills(document)
    a[0][9]=skills
    #print(df)
    csv_file_name="\\"+a[0][0]+" resume.csv"
    print(csv_file_name)
    csv_file=df.to_csv(csv_file_name,index=False)
    return(csv_file_name)

def extract_names(txt):
    person_names = []
 
    for sent in nltk.sent_tokenize(txt):
        for chunk in nltk.ne_chunk(nltk.pos_tag(nltk.word_tokenize(sent))):
            if hasattr(chunk, 'label') and chunk.label() == 'PERSON':
                person_names.append(
                    ' '.join(chunk_leave[0] for chunk_leave in chunk.leaves())
                )
 
    return person_names


def extractDataPoints(path, file_extension):
    path = str(path)
    resume_text = convert_pdf_to_txt(path)
    # print(path)
    # logger.info("Starting Data extraction from the resume- {0}".format(path))
    data_dict = {}
    
    logger = logging.getLogger(__name__)
    try:
        text = convert_pdf_to_txt(path)
        if not text:
            raise Exception
    except Exception:
        logger.exception(
            "Error extracting text from the resume- {0}".format(path))
        return {}
    try:
        clean_text = cleanText(text)
    except Exception:
        logger.exception(
            "Error performing text cleanup on the resume- {0}".format(path))
        clean_text = ''
    try:
        spacy_doc = spacyProcessText(text)
    except Exception:
        logger.exception(
            "Error processing text using Spacy on the resume- {0}".format(path))
        spacy_doc = ''
    try:
        #name = extract_full_name(spacy_doc)
        name = extract_names(resume_text)[0]
    except Exception:
        logger.exception(
            "Error extracting data point- 'Name' from the resume- {0}".format(path))
        name = ''
    try:
        skills = extract_skills(resume_text)
    except Exception:
        logger.exception(
            "Error extracting data point- 'skills from the resume- {0}".format(path))
        skills = []
    try:
        mobile_numbers = extract_phone_number(resume_text)
    except Exception:
        logger.exception(
            "Error extracting data point- 'Mobile Number' from the resume- {0}".format(path))
        mobile_numbers = []
    try:
        education = extract_education(resume_text)
    except Exception:
        logger.exception(
            "Error extracting data point- 'education from the resume- {0}".format(path))
        education = []
    try:
        emails = get_email_addresses(resume_text)
    except Exception:
        logger.exception(
            "Error extracting data point- 'Emails' from the resume- {0}".format(path))
        emails = []
    try:
        #location = extractLocation(spacy_doc)
        place_entity = locationtagger.find_locations(text =resume_text)
        countries = place_entity.other_countries
        cities = place_entity.other_regions
        if ("Tunisia" in countries) :
            location = cities[countries.index("Tunisia")] + " ,Tunisia" 
        elif ("France" in countries):
            location = cities[countries.index("France")] + " ,France" 
    except Exception:
        logger.exception(
            "Error extracting data point- 'Location' from the resume- {0}".format(path))
        location = ''

    data_dict["name"] = name
    data_dict["mobile_numbers"] = mobile_numbers
    data_dict["emails"] = emails
    data_dict["skills"] = skills
    data_dict["location"] = location 
    data_dict["education"] = education
    return data_dict


